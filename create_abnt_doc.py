from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


SOURCE = Path(r"C:\Users\chaylon\.codex\attachments\fd7f9868-861d-45fa-ade9-cca51b71a6d8\pasted-text.txt")
OUT = Path(r"C:\Users\chaylon\Desktop\ALG GUILHERME\Ponto_fixo_ABNT.docx")


def set_font(run, name="Times New Roman", size=12, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = "1"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])
    set_font(run)


def normalize_style(style, font="Times New Roman", size=12):
    style.font.name = font
    style.font.size = Pt(size)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def configure_document(doc):
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.CONTINUOUS
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = doc.styles
    normal = styles["Normal"]
    normalize_style(normal)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for name in ("Heading 1", "Heading 2"):
        style = styles[name]
        normalize_style(style)
        style.font.bold = True
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    if "EquationABNT" not in [s.name for s in styles]:
        eq = styles.add_style("EquationABNT", WD_STYLE_TYPE.PARAGRAPH)
        normalize_style(eq)
        eq.paragraph_format.first_line_indent = Cm(0)
        eq.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        eq.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        eq.paragraph_format.space_before = Pt(6)
        eq.paragraph_format.space_after = Pt(6)

    header = section.header
    header.is_linked_to_previous = False
    header.paragraphs[0].text = ""
    add_page_number(header.paragraphs[0])


def classify_line(line):
    if re.match(r"^\d+\s+[A-ZÁÉÍÓÚÂÊÔÃÕÇ0-9]", line):
        return "h1"
    if re.match(r"^\d+\.\d+\s+", line):
        return "h2"
    if (
        line.startswith("valor real =")
        or line.startswith("Δ =")
        or line.startswith("(Sx")
        or line.startswith("Sresultado =")
    ):
        return "equation"
    return "body"


def add_paragraph(doc, text, kind):
    if kind == "h1":
        p = doc.add_paragraph(style="Heading 1")
        text = text.upper()
    elif kind == "h2":
        p = doc.add_paragraph(style="Heading 2")
    elif kind == "equation":
        p = doc.add_paragraph(style="EquationABNT")
    else:
        p = doc.add_paragraph(style="Normal")

    run = p.add_run(text)
    set_font(run, bold=kind in {"h1", "h2"})
    if kind == "body":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def main():
    text = SOURCE.read_text(encoding="utf-8").replace("\ufeff", "")
    lines = [line.strip() for line in text.splitlines() if line.strip()]

    doc = Document()
    configure_document(doc)

    for line in lines:
        add_paragraph(doc, line, classify_line(line))

    doc.core_properties.title = "Representação de Números em Ponto Fixo"
    doc.core_properties.subject = "Trabalho acadêmico formatado em padrão ABNT"
    doc.core_properties.author = ""
    doc.save(OUT)
    print(str(OUT))


if __name__ == "__main__":
    main()
